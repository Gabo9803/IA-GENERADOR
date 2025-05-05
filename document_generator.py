import io
import uuid
import re
import markdown
import logging
from openai import OpenAI, AuthenticationError, RateLimitError, APIConnectionError
from cachetools import TTLCache
from utils import generate_cache_key, sanitize_fields, parse_markdown_for_pdf, summarize_history
from config import TEMPLATES, LEVEL_INSTRUCTIONS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch as reportlab_inch
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor
import os

# Configuramos el logging
logging.basicConfig(level=logging.INFO, filename='app.log', format='%(asctime)s - %(levelname)s - %(message)s')

class DocumentGenerator:
    def __init__(self, api_key: str):
        self.client = OpenAI(api_key=api_key)
        self.cache = TTLCache(maxsize=100, ttl=3600)
        self.conversation_context = {}  # Almacena el contexto por session_id

    def is_conversational_prompt(self, prompt: str) -> bool:
        """Determina si el prompt es conversacional y no requiere un documento formal."""
        conversational_keywords = [
            r'^\s*hola\s*$', r'^\s*cómo estás\s*\??\s*$', r'^\s*hey\s*$', r'^\s*hi\s*$',
            r'^\s*qué tal\s*\??\s*$', r'^\s*hello\s*$'
        ]
        prompt = prompt.lower().strip()
        return any(re.match(pattern, prompt, re.IGNORECASE) for pattern in conversational_keywords)

    def reset_context(self, session_id: str):
        """Reinicia el contexto de la conversación para una sesión específica."""
        if session_id in self.conversation_context:
            del self.conversation_context[session_id]
        logging.info(f"Contexto reiniciado para session_id: {session_id}")

    def get_prompt_suggestions(self, doc_type: str, template: str) -> list:
        """Devuelve sugerencias de prompts basadas en el tipo de documento y la plantilla."""
        suggestions = {
            'carta_formal': [
                "Redacta una carta formal invitando a un evento corporativo",
                "Escribe una carta de presentación para una solicitud de empleo",
                "Genera una carta formal de agradecimiento por una colaboración"
            ],
            'informe': [
                "Crea un informe sobre el impacto de la inteligencia artificial en la industria",
                "Redacta un informe técnico sobre energías renovables",
                "Genera un informe de progreso para un proyecto de desarrollo"
            ],
            'contrato': [
                "Escribe un contrato de prestación de servicios entre dos partes",
                "Redacta un contrato de arrendamiento para una propiedad",
                "Genera un contrato de confidencialidad para empleados"
            ],
            'factura': [
                "Crea una factura para servicios de consultoría",
                "Redacta una factura para la venta de productos",
                "Genera una factura con detalles de impuestos incluidos"
            ]
        }
        # Prompts populares generales
        popular_prompts = [
            "Redacta una carta formal invitando a un evento",
            "Escribe un informe sobre inteligencia artificial",
            "Crea un contrato de servicios profesionales",
            "Genera una factura para una venta"
        ]
        return suggestions.get(template, popular_prompts)[:3]

    def suggest_fields(self, template_content: str) -> list:
        """Sugiere campos dinámicos basados en el contenido de la plantilla."""
        # Buscar campos en formato {campo}
        fields = re.findall(r'\{(\w+)\}', template_content)
        # Añadir campos comunes si no están presentes
        common_fields = ['nombre', 'fecha', 'direccion', 'empresa']
        suggested_fields = list(set(fields + [f for f in common_fields if f not in fields]))
        return suggested_fields[:5]  # Limitar a 5 sugerencias

    def validate_generated_text(self, text: str, level: str, is_conversational: bool) -> tuple[bool, str]:
        word_count = len(text.split())
        max_words = {'basico': 500, 'medio': 1000, 'profesional': 2000}
        
        if word_count > max_words[level]:
            return False, f"El contenido excede el límite de palabras para el nivel {level} ({max_words[level]} palabras). Tiene {word_count} palabras."
        
        if word_count < 50 and not is_conversational:
            return False, "El contenido es demasiado corto (menos de 50 palabras)."
        
        if level in ['medio', 'profesional'] and not is_conversational:
            if not re.search(r'^.*\n*(# .+|## .+|### .+)', text, re.MULTILINE):
                return False, "El documento debe tener al menos un encabezado para niveles medio o profesional."
        
        return True, "Contenido válido."

    def generate(self, prompt: str, doc_type: str, template: str, fields: dict, level: str, language: str, history: list, session_id: str) -> tuple[str, bool]:
        is_conversational = self.is_conversational_prompt(prompt)
        history_summary = summarize_history(history) if not is_conversational else ""
        
        # Inicializar contexto para la sesión si no existe
        if session_id not in self.conversation_context:
            self.conversation_context[session_id] = {
                'last_document': None,
                'last_prompt': None,
                'last_doc_type': None,
                'last_template': None,
                'last_level': None,
                'last_language': None
            }

        # Añadir contexto previo al system_message si existe
        context_summary = ""
        if self.conversation_context[session_id]['last_document'] and not is_conversational:
            context_summary = (
                f"\nContexto del documento anterior:\n"
                f"- Tipo de documento: {self.conversation_context[session_id]['last_doc_type']}\n"
                f"- Plantilla: {self.conversation_context[session_id]['last_template']}\n"
                f"- Nivel: {self.conversation_context[session_id]['last_level']}\n"
                f"- Idioma: {self.conversation_context[session_id]['last_language']}\n"
                f"- Contenido previo (resumen): {self.conversation_context[session_id]['last_document'][:200]}...\n"
                "Si el usuario solicita modificaciones (por ejemplo, 'añade una cláusula'), aplica los cambios al documento anterior manteniendo su estructura y estilo."
            )

        if is_conversational:
            system_message = (
                "Eres Grok, una IA desarrollada por xAI. Responde de manera breve, amigable y directa en el idioma especificado. "
                f"Idioma: {language}. "
                "Evita generar documentos estructurados o encabezados a menos que se solicite explícitamente."
            )
        else:
            # Estructura base para todos los niveles
            system_message = (
                "Eres un asistente de IA especializado en la redacción de documentos profesionales, precisos y bien estructurados. "
                "Tu objetivo es generar contenido que sea claro, conciso y adaptado al propósito del documento. "
                f"Genera el contenido en {language}. "
                f"{LEVEL_INSTRUCTIONS[level]} "
                "Sigue estas reglas para estructurar el documento:\n"
                "- **Organización Clara y Lógica**: Organiza el contenido en secciones bien diferenciadas con subtítulos claros (#, ##, etc.). "
                "Asegúrate de que cada sección siga un flujo coherente con transiciones suaves.\n"
                "- **Introducción Ampliada**: Comienza con una **Introducción** que dé una visión general del tema, explique su contexto y su relevancia, preparando al lector para los puntos principales.\n"
                "- **Contenido Detallado y Ejemplos Prácticos**: En las secciones principales, proporciona información detallada e incluye ejemplos prácticos, casos de uso o datos ficticios realistas.\n"
                "- **Conclusión Clara y Concisa**: Termina con una **Conclusión** que resuma los puntos clave y, si corresponde, proponga futuras líneas de investigación o acción.\n"
                "- Usa un tono formal y profesional, evitando repeticiones innecesarias.\n"
                "- Usa listas con viñetas ('-') para enumerar elementos cuando sea necesario.\n"
                "- Evita jerga innecesaria y asegúrate de que el lenguaje sea accesible para un público profesional.\n"
                f"{history_summary}\n{context_summary}"
            )

            # Estructura específica según el nivel
            if level == 'basico':
                system_message += (
                    "\n**Estructura para Nivel Básico**:\n"
                    "Organiza el documento en las siguientes secciones:\n"
                    "- **# Introducción**: Proporciona una visión general del tema, su contexto y relevancia (mínimo 3-4 oraciones detalladas).\n"
                    "- **## Descripción**: Describe el tema en detalle, explicando qué es y cómo funciona (mínimo 2 párrafos).\n"
                    "- **## Conclusión**: Resume los puntos clave y menciona la importancia del tema (mínimo 2-3 oraciones).\n"
                    "Incluye al menos un ejemplo práctico simple en la sección de Descripción."
                )
            elif level == 'medio':
                system_message += (
                    "\n**Estructura para Nivel Medio**:\n"
                    "Organiza el documento en las siguientes secciones:\n"
                    "- **# Introducción**: Proporciona una visión general del tema, su contexto y relevancia.\n"
                    "- **## Descripción**: Describe el tema en detalle, explicando qué es y cómo funciona.\n"
                    "- **## Historia o Evolución**: Explica el origen o la evolución del tema a lo largo del tiempo.\n"
                    "- **## Aplicaciones y Usos Prácticos**: Detalla cómo se aplica el tema en la vida real, con ejemplos concretos.\n"
                    "- **## Conclusión**: Resume los puntos clave y propone posibles direcciones futuras.\n"
                    "Incluye ejemplos prácticos en la sección de Aplicaciones y Usos Prácticos."
                )
            elif level == 'profesional':
                system_message += (
                    "\n**Estructura para Nivel Profesional**:\n"
                    "Organiza el documento en las siguientes secciones:\n"
                    "- **# Introducción**: Proporciona una visión general del tema, su contexto y relevancia.\n"
                    "- **## Descripción**: Describe el tema en detalle, explicando qué es y cómo funciona.\n"
                    "- **## Historia o Evolución**: Explica el origen o la evolución del tema a lo largo del tiempo.\n"
                    "- **## Características Principales**: Detalla las características clave del tema.\n"
                    "- **## Aplicaciones y Usos Prácticos**: Describe aplicaciones reales, con ejemplos concretos.\n"
                    "- **## Beneficios**: Explica los beneficios del tema para los usuarios o la industria.\n"
                    "- **## Limitaciones**: Analiza las limitaciones o desafíos asociados con el tema.\n"
                    "- **## Consideraciones Éticas**: Aborda posibles riesgos éticos, como sesgos o problemas de privacidad (si aplica).\n"
                    "- **## Comparación con Alternativas**: Compara el tema con otras soluciones o tecnologías similares.\n"
                    "- **## Estudios de Caso**: Incluye un estudio de caso o ejemplo detallado de implementación.\n"
                    "- **## Impacto Futuro**: Discute cómo el tema podría evolucionar en el futuro.\n"
                    "- **## Recomendaciones**: Propón recomendaciones para su uso, implementación o mejora.\n"
                    "- **## Conclusión**: Resume los puntos clave y destaca la importancia del tema.\n"
                    "Asegúrate de incluir ejemplos prácticos, datos ficticios realistas, y análisis profundos en las secciones correspondientes."
                )

            if level in ['medio', 'profesional']:
                system_message += (
                    "\n- **Importante**: Para niveles medio y profesional, el documento debe incluir al menos un encabezado (por ejemplo, # Título, ## Subtítulo) "
                    "para estructurar el contenido de manera clara."
                )

        if template in TEMPLATES and not is_conversational:
            system_message += f" Usa esta plantilla como base:\n{TEMPLATES[template]}"

        messages = [{"role": "system", "content": system_message}] + history + [{"role": "user", "content": prompt}]
        cache_key = generate_cache_key(prompt, doc_type, template, fields, level, history)
        max_tokens = 200 if is_conversational else {'basico': 1000, 'medio': 2000, 'profesional': 4000}[level]

        if cache_key in self.cache:
            logging.info(f"Usando respuesta en caché para la clave: {cache_key}")
            return self.cache[cache_key], is_conversational

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                max_tokens=max_tokens,
                temperature=0.5
            )
            generated_text = response.choices[0].message.content.strip()
            
            is_valid, message = self.validate_generated_text(generated_text, level, is_conversational)
            if not is_valid and not is_conversational:
                logging.warning(f"Primer intento falló: {message}. Intentando regenerar con instrucciones más claras.")
                messages[-1]["content"] += "\nPor favor, asegúrate de incluir al menos un encabezado (#, ##) en el contenido y seguir la estructura solicitada según el nivel."
                response = self.client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages,
                    max_tokens=max_tokens,
                    temperature=0.5
                )
                generated_text = response.choices[0].message.content.strip()
                is_valid, message = self.validate_generated_text(generated_text, level, is_conversational)
                if not is_valid:
                    raise ValueError(f"Contenido generado no válido tras reintento: {message}")

            # Actualizar el contexto con el nuevo documento
            if not is_conversational:
                self.conversation_context[session_id].update({
                    'last_document': generated_text,
                    'last_prompt': prompt,
                    'last_doc_type': doc_type,
                    'last_template': template,
                    'last_level': level,
                    'last_language': language
                })

            self.cache[cache_key] = generated_text
            logging.info(f"Texto generado y almacenado en caché para la clave: {cache_key}")
            return generated_text, is_conversational

        except AuthenticationError:
            logging.error("Error de autenticación con la API de OpenAI.")
            raise Exception("Error de autenticación con la API de OpenAI.")
        except RateLimitError:
            logging.error("Límite de cuota alcanzado en la API de OpenAI.")
            raise Exception("Límite de cuota alcanzado. Intenta de nuevo más tarde.")
        except APIConnectionError:
            logging.error("No se pudo conectar con la API de OpenAI.")
            raise Exception("No se pudo conectar con la API de OpenAI.")
        except Exception as e:
            logging.error(f"Error inesperado al generar texto con OpenAI: {str(e)}")
            raise Exception(f"Error al generar el texto: {str(e)}")

    def extract_docx_content(self, doc):
        """Extrae el contenido de un documento DOCX como texto plano para la vista previa."""
        text = []
        section_number = 0
        subsection_number = 0

        for para in doc.paragraphs:
            if para.text.strip():
                style_name = para.style.name.lower() if para.style else 'normal'
                formatted_text = para.text.strip()

                if 'heading 1' in style_name:
                    section_number += 1
                    subsection_number = 0
                    formatted_text = f"\n\n{section_number}. {formatted_text}\n"
                elif 'heading 2' in style_name:
                    subsection_number += 1
                    formatted_text = f"\n{section_number}.{subsection_number} {formatted_text}\n"
                elif 'heading 3' in style_name:
                    formatted_text = f"{formatted_text}\n"
                elif 'list bullet' in style_name or 'list number' in style_name:
                    indent_level = para.paragraph_format.left_indent.inches / 0.5 if para.paragraph_format.left_indent else 0
                    indent = "  " * int(indent_level)
                    formatted_text = f"{indent}- {formatted_text}"
                else:
                    formatted_text = f"{formatted_text}"

                text.append(formatted_text)

        for table in doc.tables:
            text.append("\n[Tabla]\n")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text.append(" | ".join(row_text))
            text.append("\n")

        return "\n".join(text).strip()

    def parse_markdown_for_docx(self, doc, text: str, language: str, logo_path: str = None):
        """Convierte texto Markdown en un documento DOCX con estilos mejorados."""
        # Definir estilos personalizados
        styles = doc.styles

        # Estilo para Heading 1 (secciones numeradas)
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = 'Calibri'
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(79, 70, 229)  # Color #4f46e5
            h1_style.paragraph_format.space_after = Pt(12)
            h1_style.paragraph_format.space_before = Pt(18)

        # Estilo para Heading 2 (subsecciones numeradas)
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = 'Calibri'
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(79, 70, 229)
            h2_style.paragraph_format.space_after = Pt(10)
            h2_style.paragraph_format.space_before = Pt(12)

        # Estilo para Heading 3
        if 'CustomHeading3' not in styles:
            h3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
            h3_style.font.name = 'Calibri'
            h3_style.font.size = Pt(12)
            h3_style.font.bold = True
            h3_style.font.color.rgb = RGBColor(79, 70, 229)
            h3_style.paragraph_format.space_after = Pt(8)
            h3_style.paragraph_format.space_before = Pt(8)

        # Estilo para párrafos
        if 'CustomNormal' not in styles:
            normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.line_spacing = 1.15
            normal_style.paragraph_format.space_after = Pt(6)
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Estilo para listas
        if 'CustomList' not in styles:
            list_style = styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
            list_style.font.name = 'Calibri'
            list_style.font.size = Pt(11)
            list_style.paragraph_format.line_spacing = 1.15
            list_style.paragraph_format.space_after = Pt(6)
            list_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            list_style.paragraph_format.left_indent = Inches(0.5)

        # Añadir logotipo si existe
        if logo_path and os.path.exists(logo_path):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(logo_path, width=Inches(1.5))
            paragraph.paragraph_format.space_after = Pt(12)

        # Procesar el texto Markdown
        lines = text.split('\n')
        section_number = 0
        subsection_number = 0
        in_list = False
        list_level = 0
        in_table = False
        table_data = []

        for line in lines:
            line = line.strip()
            if not line:
                if in_table:
                    if table_data:
                        self._add_table_to_docx(doc, table_data)
                        table_data = []
                        in_table = False
                continue

            # Manejar encabezados
            if line.startswith('# '):
                section_number += 1
                subsection_number = 0
                section_title = f"{section_number}. {line[2:].strip()}"
                doc.add_paragraph(section_title, style='CustomHeading1')
                in_list = False
                continue
            elif line.startswith('## '):
                subsection_number += 1
                subsection_title = f"{section_number}.{subsection_number} {line[3:].strip()}"
                doc.add_paragraph(subsection_title, style='CustomHeading2')
                in_list = False
                continue
            elif line.startswith('### '):
                doc.add_paragraph(line[4:].strip(), style='CustomHeading3')
                in_list = False
                continue

            # Manejar listas
            if line.startswith('- ') or line.startswith('* '):
                if not in_list:
                    in_list = True
                    list_level = 1
                paragraph = doc.add_paragraph(line[2:].strip(), style='CustomList')
                paragraph.paragraph_format.left_indent = Inches(0.5 * list_level)
                continue
            elif line.startswith('  - ') or line.startswith('  * '):
                if not in_list:
                    in_list = True
                    list_level = 2
                else:
                    list_level = 2
                paragraph = doc.add_paragraph(line[4:].strip(), style='CustomList')
                paragraph.paragraph_format.left_indent = Inches(0.5 * list_level)
                continue
            else:
                in_list = False
                list_level = 0

            # Manejar tablas
            if line.startswith('|'):
                in_table = True
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    table_data.append(cells)
                continue
            elif in_table:
                if table_data:
                    self._add_table_to_docx(doc, table_data)
                    table_data = []
                    in_table = False

            # Párrafos normales
            doc.add_paragraph(line, style='CustomNormal')

        # Añadir la última tabla si existe
        if in_table and table_data:
            self._add_table_to_docx(doc, table_data)

    def _add_table_to_docx(self, doc, table_data):
        """Añade una tabla al documento DOCX con formato mejorado."""
        if not table_data:
            return

        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = 'Table Grid'
        table.autofit = True

        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_text
                paragraph = cell.paragraphs[0]
                paragraph.style = 'CustomNormal'
                if i == 0:  # Encabezado
                    paragraph.runs[0].font.bold = True
                    paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)
                    cell._element.get_or_add_tcPr().append(doc.element.xpath('//w:tcPr')[0].cloneNode(True))
                    shading_elm = parse_xml(r'<w:shd {} w:fill="4f46e5"/>'.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(shading_elm)

        # Ajustar el espaciado después de la tabla
        doc.add_paragraph().paragraph_format.space_before = Pt(12)

    def render(self, text: str, doc_type: str, language: str, file_name: str, logo_path: str = None) -> tuple:
        try:
            buffer = io.BytesIO()
            file_id = None
            response = text
            mime_type = None
            full_file_name = f"{file_name}.{doc_type}"
            preview_content = None

            if doc_type == 'texto':
                mime_type = 'text/plain'
                full_file_name = f"{file_name}.txt"
                buffer.write(text.encode('utf-8'))
            elif doc_type == 'markdown':
                response = markdown.markdown(text, extensions=['tables', 'fenced_code', 'nl2br'])
                mime_type = 'text/markdown'
                full_file_name = f"{file_name}.md"
                preview_content = response
                buffer.write(response.encode('utf-8'))
            elif doc_type in ['pdf', 'docx', 'html']:
                file_id = str(uuid.uuid4())
                mime_types = {
                    'pdf': 'application/pdf',
                    'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'html': 'text/html'
                }
                mime_type = mime_types[doc_type]

                if doc_type == 'pdf':
                    doc = SimpleDocTemplate(
                        buffer,
                        pagesize=letter,
                        leftMargin=0.75 * reportlab_inch,
                        rightMargin=0.75 * reportlab_inch,
                        topMargin=1 * reportlab_inch,
                        bottomMargin=1 * reportlab_inch
                    )
                    styles = getSampleStyleSheet()
                    styles['Heading1'].fontSize = 16
                    styles['Heading1'].leading = 20
                    styles['Heading1'].spaceAfter = 12
                    styles['Heading1'].fontName = 'Helvetica-Bold'
                    styles['Heading2'].fontSize = 14
                    styles['Heading2'].leading = 18
                    styles['Heading2'].spaceAfter = 10
                    styles['Heading2'].fontName = 'Helvetica-Bold'
                    styles['Heading3'].fontSize = 12
                    styles['Heading3'].leading = 16
                    styles['Heading3'].spaceAfter = 8
                    styles['Heading3'].fontName = 'Helvetica-Bold'
                    story = parse_markdown_for_pdf(text, styles, language, logo_path)
                    doc.build(story)
                    response = "PDF generado. Usa el botón de descargar para obtener el archivo."
                elif doc_type == 'docx':
                    doc = Document()
                    TWIPS_PER_INCH = 1440
                    doc.sections[0].left_margin = int(1 * TWIPS_PER_INCH)
                    doc.sections[0].right_margin = int(1 * TWIPS_PER_INCH)
                    doc.sections[0].top_margin = int(1 * TWIPS_PER_INCH)
                    doc.sections[0].bottom_margin = int(1 * TWIPS_PER_INCH)
                    self.parse_markdown_for_docx(doc, text, language, logo_path)
                    doc.save(buffer)
                    preview_content = self.extract_docx_content(doc)
                    response = "DOCX generado. Usa el botón de descargar para obtener el archivo."
                elif doc_type == 'html':
                    html_content = f"""
                    <!DOCTYPE html>
                    <html lang="{language}">
                    <head>
                        <meta charset="UTF-8">
                        <title>{file_name}</title>
                        <style>
                            body {{ 
                                font-family: 'Calibri', sans-serif; 
                                margin: 40px auto; 
                                line-height: 1.6; 
                                max-width: 900px; 
                                padding: 0 20px; 
                                color: #333; 
                                background-color: #f9f9f9; 
                            }}
                            h1 {{ 
                                color: #4f46e5; 
                                font-size: 32px; 
                                border-bottom: 3px solid #4f46e5; 
                                padding-bottom: 8px; 
                                margin-bottom: 25px; 
                                text-align: center; 
                            }}
                            h2 {{ 
                                color: #4f46e5; 
                                font-size: 24px; 
                                margin-top: 30px; 
                                margin-bottom: 15px; 
                                border-left: 5px solid #4f46e5; 
                                padding-left: 10px; 
                            }}
                            h3 {{ 
                                color: #4f46e5; 
                                font-size: 20px; 
                                margin-top: 20px; 
                                margin-bottom: 10px; 
                            }}
                            ul, ol {{ 
                                margin: 15px 0; 
                                padding-left: 30px; 
                            }}
                            li {{ 
                                margin-bottom: 10px; 
                            }}
                            table {{ 
                                border-collapse: collapse; 
                                width: 100%; 
                                margin: 20px 0; 
                                box-shadow: 0 3px 8px rgba(0,0,0,0.1); 
                                background-color: #fff; 
                            }}
                            th, td {{ 
                                border: 1px solid #ddd; 
                                padding: 12px; 
                                text-align: left; 
                            }}
                            th {{ 
                                background-color: #f0f0f0; 
                                font-weight: bold; 
                                color: #333; 
                            }}
                            p {{ 
                                margin: 12px 0; 
                                text-align: justify; 
                                font-size: 16px; 
                            }}
                            .info-box, .config-box {{ 
                                background-color: #fff; 
                                border: 2px solid #4f46e5; 
                                border-radius: 8px; 
                                padding: 20px; 
                                margin: 20px 0; 
                                box-shadow: 0 3px 8px rgba(0,0,0,0.1); 
                            }}
                            .info-box h2, .config-box h2 {{ 
                                margin-top: 0; 
                                border-left: none; 
                                padding-left: 0; 
                            }}
                            .config-box table {{ 
                                box-shadow: none; 
                                margin: 0; 
                            }}
                            hr {{ 
                                border: 0; 
                                border-top: 1px solid #ddd; 
                                margin: 20px 0; 
                            }}
                        </style>
                    </head>
                    <body>
                        <h1>{file_name}</h1>
                        <div class="info-box">
                            <h2>Información sobre la IA</h2>
                            <p>
                                Este documento fue generado por <strong>GarbotGPT</strong>, una IA desarrollada por GarolaCorp. GarBotGPT está diseñado para asistir a los usuarios en la creación de documentos profesionales y bien estructurados, ofreciendo respuestas útiles y precisas. Este documento se generó utilizando un modelo avanzado de IA con parámetros optimizados para claridad y profesionalismo.
                            </p>
                            <p>
                                <strong>Fecha de Generación:</strong> 05 de Mayo de 2025<br>
                                <strong>Idioma:</strong> {language.upper()}<br>
                                <strong>Plataforma:</strong> GarBotGPT Generador de documentos
                            </p>
                        </div>
                        <div class="config-box">
                            <h2>Configuración del Documento</h2>
                            <p>A continuación, se detalla la configuración utilizada para generar este documento:</p>
                            <table>
                                <tr>
                                    <th>Parámetro</th>
                                    <th>Valor</th>
                                </tr>
                                <tr>
                                    <td>Fuente Principal</td>
                                    <td>Calibri</td>
                                </tr>
                                <tr>
                                    <td>Tamaño de Fuente</td>
                                    <td>16px (Encabezados), 11px (Cuerpo)</td>
                                </tr>
                                <tr>
                                    <td>Color Principal</td>
                                    <td>#4f46e5</td>
                                </tr>
                                <tr>
                                    <td>Márgenes</td>
                                    <td>1 pulgada (todos los lados)</td>
                                </tr>
                                <tr>
                                    <td>Tema</td>
                                    <td>Moderno (Claro)</td>
                                </tr>
                                <tr>
                                    <td>Espaciado de Líneas</td>
                                    <td>1.15</td>
                                </tr>
                            </table>
                        </div>
                        <hr>
                        {markdown.markdown(text, extensions=['tables', 'fenced_code'])}
                    </body>
                    </html>
                    """
                    buffer.write(html_content.encode('utf-8'))
                    response = html_content
                    preview_content = html_content
                buffer.seek(0)
            else:
                logging.warning(f"Tipo de documento no soportado: {doc_type}")
                raise ValueError(f"Tipo de documento no soportado: {doc_type}")

            return response, file_id, buffer, mime_type, full_file_name, preview_content

        except Exception as e:
            logging.error(f"Error al renderizar documento: {str(e)}")
            raise Exception(f"No se pudo renderizar el documento: {str(e)}")