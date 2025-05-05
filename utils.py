import re
import json
from datetime import datetime
from hashlib import sha256
from html import escape
import io
import uuid
import matplotlib.pyplot as plt
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.platypus import KeepTogether
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from deep_translator import GoogleTranslator
import os
import logging

logging.basicConfig(level=logging.INFO, filename='app.log', format='%(asctime)s - %(levelname)s - %(message)s')

def generate_file_name(prompt: str, template: str, doc_type: str, level: str) -> str:
    words = [w for w in re.findall(r'\b\w+\b', prompt.lower()) if len(w) > 3][:3]
    base_name = '_'.join(words) if words else 'documento'
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{base_name}_{template or doc_type}_{level}_{timestamp}"

def generate_cache_key(prompt: str, doc_type: str, template: str, fields: dict, level: str, history: list) -> str:
    history_str = json.dumps(history, sort_keys=True)
    return f"{prompt}:{doc_type}:{template}:{json.dumps(fields, sort_keys=True)}:{level}:{sha256(history_str.encode()).hexdigest()}"

def sanitize_fields(fields: dict) -> dict:
    return {key: escape(str(value)) for key, value in fields.items()}

def summarize_history(history: list) -> str:
    if not history:
        return ""
    summary = "Contexto previo de la conversación:\n"
    for message in history:
        if message['role'] == 'user':
            summary += f"- El usuario dijo: {message['content'][:100]}...\n"
        elif message['role'] == 'assistant':
            summary += f"- La IA respondió: {message['content'][:100]}...\n"
    return summary

def translate_text(text: str, target_lang: str) -> str:
    if target_lang == 'es':
        return text
    try:
        translated = GoogleTranslator(source='es', target=target_lang).translate(text)
        if not translated:
            logging.warning(f"Traducción vacía para el texto '{text}' a {target_lang}, retornando texto original.")
            return text
        return translated
    except Exception as e:
        logging.error(f"Error al traducir texto '{text}' a {target_lang}: {str(e)}")
        return text

def generate_chart(data: dict, chart_type: str) -> io.BytesIO:
    buffer = io.BytesIO()
    try:
        labels = list(data.keys())
        values = list(data.values())
        if not labels or not values:
            raise ValueError("Datos de gráfico vacíos o inválidos")

        values = [float(v) for v in values]
        
        plt.figure(figsize=(6, 4))
        if chart_type == "bar":
            plt.bar(labels, values, color='skyblue')
        elif chart_type == "line":
            plt.plot(labels, values, marker='o', color='skyblue')
        plt.xlabel("Categorías")
        plt.ylabel("Valores")
        plt.title("Gráfico Generado")
        plt.grid(True)
        plt.savefig(buffer, format='png', bbox_inches='tight', dpi=100)
        plt.close()
        buffer.seek(0)
        return buffer
    except Exception as e:
        logging.error(f"Error al generar gráfico: {str(e)}")
        buffer.seek(0)
        return buffer

def add_logo_to_pdf(story: list, logo_path: str = None) -> None:
    if logo_path and os.path.exists(logo_path):
        try:
            story.append(Image(logo_path, width=100, height=50))
            story.append(Spacer(1, 20))
        except Exception as e:
            logging.error(f"Error al añadir logotipo al PDF: {str(e)}")

def add_logo_to_docx(doc: Document, logo_path: str = None) -> None:
    if logo_path and os.path.exists(logo_path):
        try:
            doc.add_picture(logo_path, width=Inches(1.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph('')
        except Exception as e:
            logging.error(f"Error al añadir logotipo al DOCX: {str(e)}")

def add_toc_to_docx(doc: Document, language: str) -> None:
    try:
        toc_title = translate_text("Índice", language)
        paragraph = doc.add_paragraph(toc_title, style='CustomTitle')
        
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar)
        
        instrText = OxmlElement('w:instrText')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        run._element.append(instrText)
        
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        run._element.append(fldChar)
        doc.add_paragraph('')
    except Exception as e:
        logging.error(f"Error al añadir tabla de contenidos al DOCX: {str(e)}")

def parse_markdown_for_pdf(text: str, styles: dict, language: str, logo_path: str = None) -> list:
    story = []
    toc = []
    toc_links = []
    lines = text.split('\n')
    in_list = False
    list_level = 0
    table_data = []
    chart_data = None
    chart_type = None

    # Estilos personalizados para PDF
    styles['Heading1'].textColor = colors.HexColor("#1a3c34")
    styles['Heading1'].fontSize = 16
    styles['Heading1'].spaceAfter = 12
    styles['Heading2'].textColor = colors.HexColor("#2e5e54")
    styles['Heading2'].fontSize = 14
    styles['Heading2'].spaceAfter = 10
    styles['Heading3'].textColor = colors.HexColor("#437f74")
    styles['Heading3'].fontSize = 12
    styles['Heading3'].spaceAfter = 8
    styles['BodyText'].fontName = 'Helvetica'
    styles['BodyText'].fontSize = 11
    styles['BodyText'].leading = 14
    styles['BodyText'].alignment = TA_JUSTIFY  # Justificar el texto

    def add_header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        header_text = translate_text("Documento Generado Automáticamente - IA Generador", language)
        footer_text = translate_text(f"Página {canvas.getPageNumber()}", language)
        canvas.drawString(72, 770, header_text)
        canvas.drawRightString(500, 50, footer_text)
        canvas.restoreState()

    add_logo_to_pdf(story, logo_path)

    # Detectar gráficos
    for line in lines:
        line = line.strip()
        if "gráfico de barras con datos:" in line.lower():
            chart_type = "bar"
            data_str = line.split("con datos:")[1].strip()
            data_pairs = data_str.split(",")
            try:
                chart_data = {pair.split(":")[0].strip(): float(pair.split(":")[1].strip()) for pair in data_pairs}
            except Exception as e:
                logging.error(f"Error al procesar datos de gráfico: {str(e)}")
                chart_data = None
            continue
        elif "gráfico de líneas con datos:" in line.lower():
            chart_type = "line"
            data_str = line.split("con datos:")[1].strip()
            data_pairs = data_str.split(",")
            try:
                chart_data = {pair.split(":")[0].strip(): float(pair.split(":")[1].strip()) for pair in data_pairs}
            except Exception as e:
                logging.error(f"Error al procesar datos de gráfico: {str(e)}")
                chart_data = None
            continue

    # Parsear el contenido
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 12))
            continue
        if line.startswith('# '):
            toc.append((line[2:], len(story), f"section_{i}"))
            story.append(Paragraph(f'<a name="section_{i}"/>', styles['BodyText']))
            story.append(Paragraph(line[2:], styles['Heading1']))
            in_list = False
        elif line.startswith('## '):
            toc.append((line[3:], len(story), f"section_{i}"))
            story.append(Paragraph(f'<a name="section_{i}"/>', styles['BodyText']))
            story.append(Paragraph(line[3:], styles['Heading2']))
            in_list = False
        elif line.startswith('### '):
            toc.append((line[4:], len(story), f"section_{i}"))
            story.append(Paragraph(f'<a name="section_{i}"/>', styles['BodyText']))
            story.append(Paragraph(line[4:], styles['Heading3']))
            in_list = False
        elif line.startswith('- ') or line.startswith('* '):
            indent = '  ' * list_level
            story.append(Paragraph(f"{indent}• {line[2:]}", styles['BodyText']))
            in_list = True
        elif line.startswith('  - ') or line.startswith('  * '):
            list_level = 1
            story.append(Paragraph(f"  ◦ {line[4:]}", styles['BodyText']))
            in_list = True
        elif line.startswith('|'):
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells and all(cells):
                table_data.append(cells)
        else:
            if table_data:
                col_widths = [100] * len(table_data[0]) if table_data else [100]
                story.append(Table(table_data, colWidths=col_widths, style=[
                    ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f0f4f0")),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ]))
                table_data = []
            if in_list:
                in_list = False
                list_level = 0
                story.append(Spacer(1, 12))
            story.append(Paragraph(line, styles['BodyText']))

    # Añadir tabla pendiente
    if table_data:
        col_widths = [100] * len(table_data[0]) if table_data else [100]
        story.append(Table(table_data, colWidths=col_widths, style=[
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#f0f4f0")),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ]))

    # Añadir gráfico si existe
    if chart_data:
        try:
            chart_buffer = generate_chart(chart_data, chart_type)
            if chart_buffer.getvalue():
                story.append(Image(chart_buffer, width=400, height=300))
                story.append(Spacer(1, 20))
            else:
                logging.warning("Buffer de gráfico vacío, no se añadió al PDF")
        except Exception as e:
            logging.error(f"Error al añadir gráfico al PDF: {str(e)}")

    # Añadir índice al inicio
    toc_title = translate_text("Índice", language)
    story.insert(0, Paragraph(toc_title, styles['Heading1']))
    for title, pos, anchor in toc:
        link = f'<link href="#{anchor}" color="blue">{title}</link>'
        story.insert(1, Paragraph(link, styles['BodyText']))
    story.insert(len(toc) + 1, Spacer(1, 20))

    return story

def parse_markdown_for_docx(doc: Document, text: str, language: str, logo_path: str = None) -> None:
    try:
        # Verificar que el texto no esté vacío
        if not text:
            raise ValueError("El texto proporcionado está vacío.")

        # Estilos personalizados para DOCX
        styles = doc.styles

        # Estilo para Título (Índice)
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', 1)
            title_style.font.name = 'Calibri'
            title_style.font.size = Pt(20)
            title_style.font.bold = True
            title_style.font.color.rgb = RGBColor(79, 70, 229)  # Color #4f46e5
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)

        # Estilo para Encabezado 1
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', 1)
            h1_style.base_style = styles['Heading 1']
            h1_style.font.name = 'Calibri'
            h1_style.font.size = Pt(16)
            h1_style.font.bold = True
            h1_style.font.color.rgb = RGBColor(79, 70, 229)
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(8)

        # Estilo para Encabezado 2
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', 1)
            h2_style.base_style = styles['Heading 2']
            h2_style.font.name = 'Calibri'
            h2_style.font.size = Pt(14)
            h2_style.font.bold = True
            h2_style.font.color.rgb = RGBColor(79, 70, 229)
            h2_style.paragraph_format.space_before = Pt(14)
            h2_style.paragraph_format.space_after = Pt(6)

        # Estilo para Párrafos
        if 'CustomNormal' not in styles:
            normal_style = styles.add_style('CustomNormal', 1)
            normal_style.font.name = 'Calibri'
            normal_style.font.size = Pt(11)
            normal_style.paragraph_format.space_after = Pt(10)
            normal_style.paragraph_format.line_spacing = 1.15
            normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Justificar el texto

        # Añadir encabezado y pie de página
        section = doc.sections[0]
        header = section.header
        header_text = translate_text("Documento Generado Automáticamente - IA Generador", language)
        header.paragraphs[0].text = header_text
        header.paragraphs[0].style.font.size = Pt(8)
        footer = section.footer
        footer_text = translate_text("Página {PAGE}", language)
        footer.paragraphs[0].text = footer_text
        footer.paragraphs[0].style.font.size = Pt(8)

        # Añadir logotipo
        add_logo_to_docx(doc, logo_path)

        # Añadir tabla de contenidos interactiva
        add_toc_to_docx(doc, language)

        # Procesar contenido
        lines = text.split('\n')
        in_list = False
        list_level = 0
        table_data = []
        chart_data = None
        chart_type = None

        # Detectar gráficos
        for line in lines:
            line = line.strip()
            if "gráfico de barras con datos:" in line.lower():
                chart_type = "bar"
                data_str = line.split("con datos:")[1].strip()
                data_pairs = data_str.split(",")
                try:
                    chart_data = {pair.split(":")[0].strip(): float(pair.split(":")[1].strip()) for pair in data_pairs}
                except Exception as e:
                    logging.error(f"Error al procesar datos de gráfico en DOCX: {str(e)}")
                    chart_data = None
                continue
            elif "gráfico de líneas con datos:" in line.lower():
                chart_type = "line"
                data_str = line.split("con datos:")[1].strip()
                data_pairs = data_str.split(",")
                try:
                    chart_data = {pair.split(":")[0].strip(): float(pair.split(":")[1].strip()) for pair in data_pairs}
                except Exception as e:
                    logging.error(f"Error al procesar datos de gráfico en DOCX: {str(e)}")
                    chart_data = None
                continue

        # Parsear el contenido
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph('')
                continue
            if line.startswith('# '):
                doc.add_paragraph(line[2:], style='CustomHeading1')
                in_list = False
            elif line.startswith('## '):
                doc.add_paragraph(line[3:], style='CustomHeading2')
                in_list = False
            elif line.startswith('### '):
                doc.add_paragraph(line[4:], style='Heading 3')
                in_list = False
            elif line.startswith('- ') or line.startswith('* '):
                style = 'List Bullet' if list_level == 0 else 'List Bullet 2'
                doc.add_paragraph(line[2:], style=style)
                in_list = True
                list_level = 0
            elif line.startswith('  - ') or line.startswith('  * '):
                list_level = 1
                doc.add_paragraph(line[4:], style='List Bullet 2')
                in_list = True
            elif line.startswith('|'):
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                if cells and all(cells):
                    table_data.append(cells)
            else:
                if table_data:
                    try:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Table Grid'
                        for i, row in enumerate(table_data):
                            for j, cell in enumerate(row):
                                table.cell(i, j).text = cell
                                table.cell(i, j).paragraphs[0].style = 'CustomNormal'
                                if i == 0:  # Estilo para el encabezado
                                    cell_run = table.cell(i, j).paragraphs[0].runs[0]
                                    cell_run.font.bold = True
                                    cell_run.font.size = Pt(11)
                                    cell_run.font.name = 'Calibri'
                                    shading_elm = OxmlElement('w:shd')
                                    shading_elm.set(qn('w:fill'), "f0f0f0")
                                    table.cell(i, j)._element.get_or_add_tcPr().append(shading_elm)
                        table_data = []
                        doc.add_paragraph('')
                    except Exception as e:
                        logging.error(f"Error al añadir tabla al DOCX: {str(e)}")
                        table_data = []
                if in_list:
                    in_list = False
                    list_level = 0
                    doc.add_paragraph('')
                paragraph = doc.add_paragraph()
                paragraph.style = 'CustomNormal'
                parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = paragraph.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*'):
                        run = paragraph.add_run(part[1:-1])
                        run.italic = True
                    else:
                        paragraph.add_run(part)

        # Añadir gráfico si existe
        if chart_data:
            try:
                chart_buffer = generate_chart(chart_data, chart_type)
                if chart_buffer.getvalue():
                    doc.add_picture(chart_buffer, width=Inches(5))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph('')
                else:
                    logging.warning("Buffer de gráfico vacío, no se añadió al DOCX")
            except Exception as e:
                logging.error(f"Error al añadir gráfico al DOCX: {str(e)}")

    except Exception as e:
        logging.error(f"Error al generar DOCX en parse_markdown_for_docx: {str(e)}")
        raise Exception(f"No se pudo generar el documento DOCX: {str(e)}")